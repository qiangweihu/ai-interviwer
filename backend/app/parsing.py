from __future__ import annotations

import re
import tempfile
from pathlib import Path


class ResumeParseError(ValueError):
    pass


EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+(?:\.[\w-]+)+\b")
PHONE_RE = re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)(?!\d)")


def _redact_contacts(text: str) -> str:
    text = EMAIL_RE.sub("[已移除邮箱]", text)
    return PHONE_RE.sub("[已移除电话]", text)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ResumeParseError("PDF 已加密，无法读取。请上传未加密的文本版 PDF。")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def parse_resume(filename: str, content: bytes, max_bytes: int) -> str:
    """Parse a supported resume while ensuring the temporary source is removed.

    The returned text is intentionally transient: callers should pass it to the
    profile extractor and persist only the resulting structured summary.
    """

    if len(content) > max_bytes:
        raise ResumeParseError(f"简历超过 {max_bytes // (1024 * 1024)} MB 限制。")
    suffix = Path(filename or "resume.txt").suffix.lower()
    supported = {".pdf", ".docx", ".md", ".markdown", ".txt"}
    if suffix not in supported:
        raise ResumeParseError("仅支持 PDF、DOCX、Markdown 或 TXT 简历。")

    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(prefix="resume-", suffix=suffix, delete=False) as handle:
            handle.write(content)
            temp_path = Path(handle.name)
        if suffix == ".pdf":
            text = _read_pdf(temp_path)
        elif suffix == ".docx":
            text = _read_docx(temp_path)
        else:
            text = content.decode("utf-8-sig", errors="replace")
        text = _redact_contacts(text).strip()
        if not text:
            raise ResumeParseError("未读取到简历文字。请上传带文本层的 PDF 或可读文档。")
        return text
    except ResumeParseError:
        raise
    except Exception as exc:  # pragma: no cover - parser library-specific errors
        raise ResumeParseError(f"简历读取失败：{exc}") from exc
    finally:
        if temp_path:
            try:
                temp_path.unlink(missing_ok=True)
            except OSError:
                pass
